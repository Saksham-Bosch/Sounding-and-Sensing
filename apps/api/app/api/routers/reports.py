import io
import os

import docx
import httpx
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse

from app.repositories.excel_adapter import ExcelDatabase

router = APIRouter(prefix="/reports", tags=["Reports"])
db = ExcelDatabase()

AZURE_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_MODEL = os.getenv("AZURE_OPENAI_CHAT_MODEL")
AZURE_VERSION = os.getenv("AZURE_OPENAI_API_VERSION")


async def generate_ai_summary(event_data: dict, answers_data: list, audience: str) -> str:
    """Uses Azure OpenAI to cluster responses into an audience-specific summary."""
    if not all([AZURE_ENDPOINT, AZURE_API_KEY, AZURE_MODEL]):
        return "[Summary generation skipped: Azure credentials missing]"

    # Build the context from the event and all participant answers
    context = (
        f"Event: {event_data.get('title')}\n"
        f"Dates: {event_data.get('start_date')} to {event_data.get('end_date')}\n"
        f"Speakers: {', '.join(event_data.get('speakers', []))}\n"
        f"Description: {event_data.get('description')}\n\n"
        f"Participant Responses:\n"
    )
    for ans in answers_data:
        context += f"- [Role: {ans.get('participant_role', 'Unknown')}] Question: {ans.get('question_text')} | Answer: {ans.get('normalized_text')}\n"

    system_prompt = (
        f"You are an expert business analyst. Synthesize the provided participant responses into an executive summary "
        f"tailored specifically for a '{audience}' audience. Group similar themes, highlight key metrics, and extract actionable insights. "
        f"Use clear markdown headings and bullet points."
    )

    url = f"{AZURE_ENDPOINT}/openai/deployments/{AZURE_MODEL}/chat/completions?api-version={AZURE_VERSION}"
    headers = {"api-key": AZURE_API_KEY, "Content-Type": "application/json"}
    payload = {
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": context}
        ],
        "max_completion_tokens": 1500
    }

    async with httpx.AsyncClient(timeout=60.0, verify=False) as client:
        try:
            response = await client.post(url, headers=headers, json=payload)
            response.raise_for_status()
            return response.json()["choices"][0]["message"]["content"].strip()
        except Exception as e:
            return f"Failed to generate summary: {str(e)}"


@router.get("/events/{event_id}/summary")
async def get_event_summary(event_id: str, audience_level: str = "Executive Board"):
    """Generates an AI summary of all answers for an event."""
    events = db.get_all_records("Events")
    event = next((e for e in events if e.get("id") == event_id), None)
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")

    # Map session IDs to participant roles
    sessions = {s["session_id"]: s.get("participant_role", "General") for s in db.get_all_records("InterviewSessions") if s.get("event_id") == event_id}

    # Map question IDs to question text
    questions = {q["id"]: q["text"] for q in db.get_all_records("GeneratedQuestions")}

    # Filter answers that belong to this event's sessions
    event_answers = [
        {
            "participant_role": sessions[ans["session_id"]],
            "question_text": questions.get(ans.get("question_id"), "Unknown Question"),
            "normalized_text": ans.get("normalized_text", "")
        }
        for ans in db.get_all_records("InterviewAnswers") if ans.get("session_id") in sessions
    ]

    if not event_answers:
        return {"audience": audience_level, "summary": "No participant answers found for this event yet."}

    summary_text = await generate_ai_summary(event, event_answers, audience_level)
    return {"audience": audience_level, "summary": summary_text}


@router.get("/events/{event_id}/export")
async def export_event_report_docx(event_id: str, audience_level: str = "Executive Board"):
    """Generates a formatted DOCX report based on the AI Summary."""
    events = db.get_all_records("Events")
    event = next((e for e in events if e.get("id") == event_id), None)
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")

    # 1. Get the generated summary for this specific audience
    summary_data = await get_event_summary(event_id, audience_level)
    summary_text = summary_data.get("summary", "No data.")

    # 2. Build the Word Document
    doc = docx.Document()
    doc.add_heading(f"Event Report: {event.get('title', 'Unknown')}", 0)
    doc.add_paragraph(f"Audience Perspective: {audience_level}")

    doc.add_heading("Event Description", level=2)
    doc.add_paragraph(event.get("description", ""))

    doc.add_heading(f"AI Aggregated Summary ({audience_level})", level=2)

    # Simple markdown parsing for the DOCX
    for line in summary_text.split("\n"):
        if line.startswith("#"):
            level = line.count("#")
            doc.add_heading(line.replace("#", "").strip(), level=min(level + 1, 9))
        elif line.strip():
            doc.add_paragraph(line.strip())

    # 3. Save to in-memory stream and return as downloadable file
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"Report_{event.get('title', 'Event')}_{audience_level}.docx".replace(" ", "_")
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
